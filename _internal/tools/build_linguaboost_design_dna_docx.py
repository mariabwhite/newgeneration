from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT.parent
OUT_DIR = ROOT / "prompt-library" / "technical-prompts"
OUT = OUT_DIR / "linguaboost-lab-design-dna-technical-prompt.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D4C8EA", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=120, bottom=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_run(run, font="Manrope", size=10.5, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, size=10.5, color="1A1326", bold=False, space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    style_run(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(text)
    if level == 1:
        style_run(r, font="Unbounded", size=18, bold=True, color="4F3EA5")
    elif level == 2:
        style_run(r, font="Unbounded", size=13, bold=True, color="7A5FCF")
    else:
        style_run(r, font="Manrope", size=11, bold=True, color="1A1326")
    return p


def add_code(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3EEFF")
    set_cell_border(cell, "C7B6EA", "6")
    set_cell_margins(cell, 140, 160, 140, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        style_run(r, font="JetBrains Mono", size=8.2, color="1A1326")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "E8E0F0")
        set_cell_border(hdr[i], "B8A6E3", "8")
        set_cell_margins(hdr[i])
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        style_run(r, font="Manrope", size=9.2, bold=True, color="1A1326")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_border(cells[i], "D6CBEA", "4")
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            r = p.add_run(str(value))
            font = "JetBrains Mono" if any(ch in str(value) for ch in "#():;%") else "Manrope"
            style_run(r, font=font, size=8.7, color="1A1326")
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        r = p.add_run(item)
        style_run(r, size=9.8, color="1A1326")


def add_image_if_exists(doc, path, caption, width_inches=6.8):
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_inches))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        style_run(r, size=8.5, color="5F5A70")


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.45)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)

    styles = doc.styles
    styles["Normal"].font.name = "Manrope"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Manrope")
    styles["Normal"].font.size = Pt(10.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("LinguaBoost Lab")
    style_run(r, font="Unbounded", size=24, bold=True, color="4F3EA5")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("письменный слепок ДНК дизайна и технический промпт")
    style_run(r, font="Manrope", size=13, bold=True, color="FF5A1F")

    meta = [
        ("Источник", "lingua-boost-lab/index.html"),
        ("Статус", "Снято после последних ручных правок: мягкие края light-баннера, gap 1cm, фоновые индексы 00."),
        ("Назначение", "Передать другой модели или разработчику полный визуальный контракт страницы лаборатории."),
        ("Папка", "_internal/prompt-library/technical-prompts"),
    ]
    add_table(doc, ["Поле", "Значение"], meta, [3.2, 12.0])

    add_heading(doc, "1. Короткая формула ДНК", 1)
    add_para(
        doc,
        "Это интерфейс учебной лаборатории LinguaBoost: плотный, технологичный, но не SaaS-сухой. "
        "Основа - тёмная plum-тема с оранжевыми CTA и светлая lavender-тема. Hero строится как "
        "полноширинная сцена с текстом слева и фотореалистичным учебным сетапом справа. Каталог ниже "
        "работает как строгая сетка модулей: уровни, прогресс, карточки уроков, терминальный микроязык.",
    )
    add_bullets(
        doc,
        [
            "Главная эмоция: умная учебная лаборатория, не лендинг и не декоративная открытка.",
            "Композиция: left copy + right hero art; каталог начинается через 1cm после уровня-кнопок.",
            "Графический код: plum, violet, orange, monospace labels, dashed dividers, ghost-index 00.",
            "Форма UI: радиус 8-10px, тонкие borders, минимум теней, чёткие кнопки и плотная навигация.",
        ],
    )

    add_heading(doc, "2. Цветовая система", 1)
    dark_rows = [
        ("--bg", "#0e0a18", "основной тёмный фон"),
        ("--bg-2", "#120c1d", "второй слой фона"),
        ("--bg-card", "#1a1326", "карточки"),
        ("--bg-card-2", "#21172f", "приподнятые панели"),
        ("--plum", "#9a7ad9", "главный фиолетовый"),
        ("--plum-2", "#c5aef0", "светлый фиолетовый"),
        ("--accent", "#ff9b2f", "оранжевый CTA / terminal marker"),
        ("--success", "#2ee59d", "зелёный прогресс"),
        ("--text", "#f0ebff", "основной текст"),
        ("--muted", "#bfaee0", "вторичный текст"),
        ("--line", "rgba(197,174,240,0.12)", "тонкие границы"),
        ("--line-2", "rgba(197,174,240,0.22)", "активные границы"),
        ("--module-bg", "rgba(12,8,20,0.72)", "фон модулей"),
    ]
    light_rows = [
        ("--bg", "#f7f5ff", "исходный light-фон"),
        ("edge-match", "#e8e0f0", "финальный фон страницы и hero"),
        ("--bg-card", "#ffffff", "база карточек"),
        ("--bg-card-2", "#f2eeff", "приподнятые панели"),
        ("--text", "#1a1326", "основной текст"),
        ("--muted", "#5f5a70", "вторичный текст"),
        ("--plum", "#7a5fcf", "главный фиолетовый"),
        ("--plum-2", "#4f3ea5", "глубокий фиолетовый"),
        ("--accent", "#ff5a1f", "оранжевый CTA"),
        ("--line", "rgba(122,95,207,0.18)", "тонкие границы"),
        ("--line-dot", "rgba(122,95,207,0.42)", "пунктир"),
        ("module card", "rgba(255,255,255,0.46)", "финальный фон карточек"),
    ]
    add_heading(doc, "Dark palette", 2)
    add_table(doc, ["Token", "Value", "Use"], dark_rows, [3.6, 4.4, 7.2])
    add_heading(doc, "Light palette", 2)
    add_table(doc, ["Token", "Value", "Use"], light_rows, [3.6, 4.4, 7.2])

    add_heading(doc, "3. Типографика", 1)
    add_table(
        doc,
        ["Элемент", "Шрифт", "Размер / вес / трекинг", "Поведение"],
        [
            ("Display H1", "Unbounded", "clamp(2.2rem, 4.2vw, 3.8rem), 800, -0.06em", "line-height 0.94; плотный заголовок"),
            ("Hero final H1", "Unbounded", "max-width min(860px, 60vw)", "акценты через .accent-purple"),
            ("Body / lead", "Manrope", "15-16px, 400/500", "max-width min(680px, 46vw)"),
            ("Mono labels", "JetBrains Mono", "10-13px, 700/800, uppercase", "topbar, badges, level count, pills"),
            ("Level H2", "Unbounded", "clamp(64px, 8.5vw, 100px), 900", "line-height 0.88, letter-spacing -0.08em"),
            ("Ghost index", "Unbounded", "clamp(74px, 7.36vw, 110px), 900", "line-height 0.72, letter-spacing -0.10em"),
        ],
        [3.5, 3.2, 5.5, 5.0],
    )

    add_heading(doc, "4. Layout и размеры", 1)
    add_table(
        doc,
        ["Блок", "Финальные размеры / отступы", "Комментарий"],
        [
            ("topbar", "height 64px; padding 0 28px; sticky top 0; z-index 100", "размытый стеклянный бар"),
            ("brand-logo", "height 38px; max-width min(218px, 34vw)", "dark/light logo swap"),
            ("wrap", "max-width 1320px; margin auto; base padding 48px 28px 80px", "контейнер страницы"),
            ("hero", "width 100vw; margin-top calc(-64px - 34px); padding-top calc(120px + 34px - 2cm)", "сцена уходит под topbar"),
            ("hero bottom", "margin-bottom 1cm; hero-copy padding-bottom 0", "каталог начинается близко к row уровней"),
            ("hero art", "right 0; width min(44.7vw, 847px); top calc(50% + 2cm)", "оба ассета опущены"),
            ("level-grid", "padding-top 0; padding-left/right historically 4cm in early base", "финальный верх убран"),
            ("level-section", "padding 38px 36px 36px; border-radius 10-12px", "карточка уровня"),
            ("module-grid", "repeat(8, minmax(0, 1fr)); gap 12px", "одна строка из 8 уроков"),
            ("module-card", "min-height 138px; padding 18px 12px; gap 10px; radius 10px", "компактные карточки"),
            ("module-link", "min-height 36px; padding 0 16px; radius 8px", "оранжевые кнопки"),
        ],
        [3.2, 7.2, 5.0],
    )

    add_heading(doc, "5. Hero DNA", 1)
    add_bullets(
        doc,
        [
            "Hero не должен выглядеть как карточка: border, radius и shadow убраны; сцена полноширинная.",
            "Изображение справа является главным визуальным активом, но не перекрывает текст; оно маскируется слева.",
            "В light-теме фон выровнен по плоскости картинки: #e8e0f0. Это убирает прямоугольную заплатку.",
            "Light-арт не blurится целиком. Смягчаются только края через mask-image с несколькими stop-точками.",
            "Текстовая колонка смещена translateX(-38px), max-width 1240px, padding 30px 32px 0.",
        ],
    )
    add_code(
        doc,
        ".hero { width: 100vw; margin-left: calc(50% - 50vw); margin-top: calc(-64px - 34px); }\n"
        ".hero { padding-top: calc(120px + 34px - 2cm); margin-bottom: 1cm; }\n"
        ".hero-banner img.hero-bg { right: 0; width: min(44.7vw, 847px); top: calc(50% + 2cm); }\n"
        ".hero-copy { max-width: 1240px; padding: 30px 32px 0; transform: translateX(-38px); }"
    )

    add_heading(doc, "6. Catalog DNA", 1)
    add_bullets(
        doc,
        [
            "Каталог начинается почти сразу после row уровней: между controls и первой секцией держать 1cm.",
            "Секция уровня имеет крупный H2 слева, terminal-count справа, пунктирную линию под header.",
            "Ghost-index 00 должен быть фоновым, не главным: справа, над прогресс-панелью, прозрачный.",
            "Финальное положение 00: top calc(84px - 1cm); right 40px; font clamp(74px, 7.36vw, 110px).",
            "Пунктир section-head: border-top 1px dashed rgba(197,174,240,0.34), light rgba(122,95,207,0.30).",
        ],
    )
    add_code(
        doc,
        ".level-section::after {\n"
        "  top: calc(84px - 1cm); right: 40px;\n"
        "  font-size: clamp(74px, 7.36vw, 110px);\n"
        "  line-height: 0.72; letter-spacing: -0.10em;\n"
        "  color: rgba(154, 122, 217, 0.16);\n"
        "}\n"
        "[data-theme='light'] .level-section::after { color: rgba(122, 95, 207, 0.13); }"
    )

    add_heading(doc, "7. Components", 1)
    add_table(
        doc,
        ["Компонент", "Техническая ДНК", "Визуальная роль"],
        [
            ("Level buttons", "padding 8px 14px; radius 10px; mono 12px; active border accent", "быстрая навигация по уровням"),
            ("Hero CTA primary", "orange gradient/solid; white or dark text by theme; mono uppercase", "главное действие"),
            ("Hero CTA secondary", "transparent; line border; muted text", "тихий переход к программам"),
            ("Terminal level-count", "width 220px; mono 12px; border line-2; translateY(-18px)", "служебный статус уровня"),
            ("Progress panel", "padding 20px 22px; radius 8-9px; track radius 999px", "живой прогресс A1"),
            ("Module pill", "32px min-height; dotted border; mono 12px", "уровень и номер урока"),
            ("Module card", "min-height 138px; no shadow; border line-2", "плотная сетка уроков"),
            ("Soon/disabled link", "light: rgba(122,95,207,0.08); muted violet", "неактивные уроки без оранжевого шума"),
        ],
        [3.4, 6.6, 5.2],
    )

    add_heading(doc, "8. Assets", 1)
    add_table(
        doc,
        ["Asset", "Path", "Правило"],
        [
            ("Dark logo", "lingua-boost-lab/assets/linguaboost-logo-dark-lockup.png", "виден в dark"),
            ("Light logo", "lingua-boost-lab/assets/linguaboost-logo-light-lockup-purple.png", "виден в light"),
            ("Dark hero", "lingua-boost-lab/assets/lab/lab-hero-dark.png", "dark hero art"),
            ("Light hero", "lingua-boost-lab/assets/lab/lab-hero-light.png", "light hero art with feathered mask"),
            ("Favicon", "assets/linguaboost-lab-favicon.svg", "page icon"),
            ("Drop icons", "linguaboost-drop-purple.svg / orange.svg", "brand mark fallback"),
        ],
        [3.2, 7.2, 5.0],
    )

    add_heading(doc, "9. Технический промпт для библиотеки", 1)
    prompt = (
        "Задача: восстановить или продолжить страницу LinguaBoost Lab в точном стиле текущего дизайна. "
        "Работай как senior frontend designer/developer. Сохрани два режима темы: dark Midnight Plum и light Lavender. "
        "Не делай landing page, не добавляй декоративные карточки вокруг hero. Первый экран должен быть рабочей лабораторией: "
        "sticky topbar 64px, полноширинный hero, слева terminal badge '$ lab --level', крупный Unbounded H1, lead, два CTA, "
        "ряд уровней PRE-A1/A1/A2/B1/B2/C1/C2/Interactive. Справа hero image: dark/light PNG из assets/lab, width min(44.7vw, 847px), "
        "right 0, top calc(50% + 2cm), object-fit contain. В light теме не размывай сам арт; смягчай только края mask-image, "
        "фон hero и страницы #e8e0f0. Между hero controls и каталогом держи gap 1cm: hero margin-bottom 1cm, hero-copy padding-bottom 0, "
        "level-grid padding-top 0. Каталог: level-section с radius 10-12px, тонкая border, H2 level слева, terminal level-count справа, "
        "пунктирная линия под header. Фоновый индекс 00: top calc(84px - 1cm), right 40px, font-size clamp(74px, 7.36vw, 110px), "
        "line-height 0.72, letter-spacing -0.10em, opacity через rgba. Module grid: 8 columns, gap 12px. Cards: min-height 138px, "
        "padding 18px 12px, no shadow, border line-2. CTA/link buttons: orange #ff7430 -> #ff5125, radius 8px, mono uppercase. "
        "Не используй большие marketing hero, gradient blobs, nested cards, oversized round pills. Проверяй desktop и mobile screenshots."
    )
    add_code(doc, prompt)

    add_heading(doc, "10. Финальные override-правила, которые нельзя потерять", 1)
    add_code(
        doc,
        "[data-theme='light'], body.light, body.light .hero, body.light .level-grid { background: #e8e0f0 !important; }\n"
        "[data-theme='light'] .hero-banner img.hero-bg { filter: saturate(1.08) contrast(1.08) brightness(1) !important; }\n"
        ".hero { padding-top: calc(120px + 34px - 2cm) !important; margin-bottom: 1cm !important; }\n"
        ".hero-copy { padding-bottom: 0 !important; }\n"
        ".level-grid { padding-top: 0 !important; }\n"
        ".hero-banner img.hero-bg { top: calc(50% + 2cm) !important; }\n"
        ".level-section::after { top: calc(84px - 1cm) !important; font-size: clamp(74px, 7.36vw, 110px) !important; }"
    )

    add_heading(doc, "11. Визуальные контрольные снимки", 1)
    add_image_if_exists(
        doc,
        PROJECT / "audit-shots" / "codex-lab-gap-1cm-light-v2.png",
        "Light theme: hero, softened edges, catalog begins after 1cm gap.",
        6.9,
    )
    add_image_if_exists(
        doc,
        PROJECT / "audit-shots" / "codex-lab-gap-1cm-dark.png",
        "Dark theme: same layout rhythm and catalog spacing.",
        6.9,
    )

    add_heading(doc, "12. QA checklist", 1)
    add_bullets(
        doc,
        [
            "Hero art visible and sharp in both themes; light edges feathered, not globally blurred.",
            "Catalog top starts approximately 1cm after the level selector row.",
            "Ghost 00 is visible but quiet, sits near section header/progress zone without dominating.",
            "Module grid remains 8 columns on desktop; responsive rules may collapse below 900px and 600px.",
            "No visible orange glow leaks into light module cards except true CTA buttons.",
            "Typography remains Unbounded + Manrope + JetBrains Mono; no default browser fonts.",
        ],
    )

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    add_heading(doc, "Appendix: raw token map", 1)
    add_table(
        doc,
        ["Group", "Values"],
        [
            ("Fonts", "Unbounded 700/800/900; Manrope 400/500/700/800; JetBrains Mono 400/700"),
            ("Radii", "5px pills, 8px controls, 9px progress, 10px cards/buttons, 12px legacy sections"),
            ("Borders", "1px solid rgba lines; dashed divider; dotted module pills"),
            ("Motion", "hover transitions 0.16-0.18s; no heavy animation on index page"),
            ("Blend/mask", "hero light mix-blend multiply earlier, final feather mask; dark left fade mask"),
            ("Spacing units", "cm intentionally used for hero/catalog micro-adjustments: 1cm, 2cm"),
        ],
        [3.0, 12.2],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
